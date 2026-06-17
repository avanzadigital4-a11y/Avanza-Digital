# -*- coding: utf-8 -*-
"""
jarvis_contratos.py — Generador de contratos de prestación de servicios (PDF).

Mismo mecanismo que jarvis_propuestas.py: toma datos (cliente + plan + datos
fiscales) y devuelve un documento armado. Acá el documento final es un PDF
renderizado con xhtml2pdf (pisa) a partir de una plantilla HTML+CSS.

Uso típico desde una ruta FastAPI (ver jarvis_contratos_routes.py):

    import jarvis_contratos as contratos
    datos = contratos.datos_desde_venta(venta, extra={...campos fiscales...})
    pdf_bytes = contratos.render_contrato_pdf(datos)

El Anexo I (entregables) se autocompleta desde PLAN_INFO, que replica
CONTENIDOS_POR_PLAN / PLANES_AVANZA de jarvis_propuestas.py. Si más adelante
querés una única fuente de verdad, mové PLAN_INFO a jarvis_config.py y
que ambos módulos lo importen de ahí.

Nota: el texto legal es una base profesional pero NO sustituye la revisión
de un abogado matriculado. Revisalo antes de usarlo en producción.
"""
from __future__ import annotations

import os
import datetime
import unicodedata
from dataclasses import dataclass, field
from typing import Optional

# ──────────────────────────────────────────────────────────────────────────────
# 1) DATOS DE LA EMPRESA (constantes) — editá una sola vez acá o por variables
#    de entorno. Idealmente esto va a jarvis_config.py.
# ──────────────────────────────────────────────────────────────────────────────
AVANZA = {
    "razon_social": os.environ.get("AVANZA_RAZON_SOCIAL", "Avanza Digital"),
    "cuit":         os.environ.get("AVANZA_CUIT", "[CUIT de Avanza]"),
    "domicilio":    os.environ.get("AVANZA_DOMICILIO", "Santa Fe, Argentina"),
    "representante":os.environ.get("AVANZA_REPRESENTANTE", "[representante de Avanza]"),
    "cargo":        os.environ.get("AVANZA_CARGO", "Titular"),
    "email":        os.environ.get("AVANZA_EMAIL", "contacto@avanzadigital.digital"),
    "jurisdiccion": os.environ.get("AVANZA_JURISDICCION", "Santa Fe"),
    "web":          "avanzadigital.digital",
    # Logo opcional: ruta a un PNG/SVG local o data-URI. Si no existe, se usa
    # un wordmark tipográfico.
    "logo_path":    os.environ.get("AVANZA_LOGO_PATH", ""),
}

# ──────────────────────────────────────────────────────────────────────────────
# 2) INFO POR PLAN — precio, plazo, garantía, capacitación y entregables.
#    (Replica jarvis_propuestas.PLANES_AVANZA + CONTENIDOS_POR_PLAN.)
# ──────────────────────────────────────────────────────────────────────────────
PLAN_INFO = {
    "Plan Base": {
        "precio": 1050.0,
        "plazo_dias": 7,
        "garantia": "90 (noventa) días corridos",
        "capacitacion": 0,
        "sla": False,
        "entregables": [
            "Sitio web profesional hasta 5 páginas",
            "Diseño adaptado al rubro",
            "Optimización SEO básica",
            "Formulario de contacto y WhatsApp integrado",
            "Alta en Google Maps y Google Business",
            "Panel de administración simple",
        ],
    },
    "Plan Pro": {
        "precio": 2900.0,
        "plazo_dias": 14,
        "garantia": "90 (noventa) días corridos",
        "capacitacion": 1,
        "sla": False,
        "entregables": [
            "Sitio web profesional ilimitado",
            "Diseño premium adaptado al sector industrial",
            "SEO avanzado con posicionamiento local y sectorial",
            "Sistema de captación de leads con CRM básico",
            "Integración WhatsApp Business con respuesta automática",
            "Formularios inteligentes de cotización",
            "Alta y gestión de Google Business Profile",
            "Informe mensual de performance y leads generados",
            "Soporte prioritario con SLA de 24 hs",
        ],
    },
    "Plan Industrial": {
        "precio": 4900.0,
        "plazo_dias": 21,
        "garantia": "90 (noventa) días corridos",
        "capacitacion": 1,
        "sla": True,
        "entregables": [
            "Todo el Plan Pro incluido",
            "Catálogo de productos/servicios con fichas técnicas",
            "Sistema de cotización online automatizado",
            "Integración con CRM propio del cliente",
            "SEO industrial avanzado — posicionamiento por producto y proceso",
            "Landings específicas por línea de producto o servicio",
            "Campañas de Google Ads industriales básicas",
            "Informe quincenal con análisis de competencia",
            "Consultoría estratégica mensual (1 h)",
            "IA JARVIS para el equipo comercial incluida",
            "Panel de métricas en tiempo real",
        ],
    },
    "Estrategico 360": {
        "precio": 7500.0,
        "plazo_dias": 30,
        "garantia": "12 (doce) meses",   # garantía extendida prometida en el brochure
        "capacitacion": 2,
        "sla": True,
        "entregables": [
            "Todo el Plan Industrial incluido",
            "Desarrollo a medida / intranet",
            "Integración con ERP (Tango, SAP) — bidireccional",
            "Automatización condicional y Lead Scoring",
            "Estrategia digital integral a 12 meses",
            "Campañas de Google Ads y LinkedIn Ads industriales",
            "Automatización de marketing y nurturing de leads",
            "IA JARVIS completa con módulos avanzados",
            "Consultoría estratégica quincenal (2 h)",
            "Gerente de cuenta dedicado",
            "Reporte ejecutivo mensual con ROI medido",
            "Garantía extendida (12 meses) y Soporte con SLA",
        ],
    },
}

# Aliases tolerantes para el campo `plan` (que puede venir con/sin acento, etc.)
_ALIASES = {
    "base": "Plan Base",
    "planbase": "Plan Base",
    "pro": "Plan Pro",
    "planpro": "Plan Pro",
    "industrial": "Plan Industrial",
    "planindustrial": "Plan Industrial",
    "estrategico360": "Estrategico 360",
    "estrategico": "Estrategico 360",
    "planestrategico360": "Estrategico 360",
    "estrategico360grados": "Estrategico 360",
}

MESES_ES = ["", "enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]


def _strip_accents(s: str) -> str:
    return "".join(c for c in unicodedata.normalize("NFD", s)
                   if unicodedata.category(c) != "Mn")


def normalizar_plan(plan: Optional[str]) -> str:
    """Mapea cualquier variante del nombre del plan a la clave canónica."""
    if not plan:
        return "Plan Pro"
    key = _strip_accents(plan).lower().replace(" ", "").replace("-", "")
    key = key.replace("°", "").replace("º", "")
    if key in _ALIASES:
        return _ALIASES[key]
    # match laxo por substring
    for alias, canon in _ALIASES.items():
        if alias in key:
            return canon
    return plan if plan in PLAN_INFO else "Plan Pro"


# ──────────────────────────────────────────────────────────────────────────────
# 3.bis) TERMINOLOGÍA FISCAL / IDENTIFICATORIA POR PAÍS
#   Resuelve el "¿qué es el CUIT?": cada país nombra distinto el identificador
#   tributario de una EMPRESA y el documento de una PERSONA. El contrato muestra
#   el término local del CLIENTE; AVANZA (prestador argentino) sigue usando CUIT.
# ──────────────────────────────────────────────────────────────────────────────
IDENT_FISCAL_POR_PAIS = {
    "Argentina":  {"empresa": "CUIT",            "persona": "DNI",       "ej": "30-71234567-8"},
    "Mexico":     {"empresa": "RFC",             "persona": "CURP / INE","ej": "ABC120101AB1"},
    "Peru":       {"empresa": "RUC",             "persona": "DNI",       "ej": "20123456789"},
    "Chile":      {"empresa": "RUT",             "persona": "RUN",       "ej": "76.123.456-7"},
    "Colombia":   {"empresa": "NIT",             "persona": "Cédula",    "ej": "900.123.456-7"},
    "Costa Rica": {"empresa": "Cédula jurídica", "persona": "Cédula",    "ej": "3-101-123456"},
    "Venezuela":  {"empresa": "RIF",             "persona": "Cédula",    "ej": "J-12345678-9"},
    "Uruguay":    {"empresa": "RUT",             "persona": "C.I.",      "ej": "212345670019"},
    "Paraguay":   {"empresa": "RUC",             "persona": "C.I.",      "ej": "80012345-6"},
    "Ecuador":    {"empresa": "RUC",             "persona": "Cédula",    "ej": "1790012345001"},
    "Bolivia":    {"empresa": "NIT",             "persona": "C.I.",      "ej": "1234567890"},
    "España":     {"empresa": "CIF / NIF",       "persona": "DNI / NIE", "ej": "B12345678"},
}
_IDENT_DEFAULT = {"empresa": "identificación tributaria",
                  "persona": "documento de identidad", "ej": ""}


def ident_fiscal(pais):
    """Términos fiscales/identificatorios locales para un país.
    Tolerante a acentos y mayúsculas. Si no se reconoce el país, cae a un
    término genérico ("identificación tributaria" / "documento de identidad")."""
    if not pais:
        return dict(IDENT_FISCAL_POR_PAIS["Argentina"])
    clave = _strip_accents(pais).strip().lower()
    for nombre, datos in IDENT_FISCAL_POR_PAIS.items():
        if _strip_accents(nombre).lower() == clave:
            return dict(datos)
    return dict(_IDENT_DEFAULT)


# ──────────────────────────────────────────────────────────────────────────────
# 3.ter) PLANES DE MANTENIMIENTO (OPCIONALES) — espejo de PLANES_CONTINUIDAD
#   Sirven para el Anexo II. El mantenimiento NUNCA es obligatorio; solo aparece
#   en el contrato si el aliado lo marca (incluir_mantenimiento=True).
# ──────────────────────────────────────────────────────────────────────────────
MANTENIMIENTO_INFO = {
    "Plan Cuidado": {"precio": 80.0, "items": [
        "Hosting profesional de alta velocidad",
        "Dominio profesional (.com, .com.ar o local) y Certificado SSL",
        "Backups automáticos semanales",
        "Seguridad y monitoreo 24/7",
        "Soporte técnico por fallas del sistema",
        "Reporte básico mensual"]},
    "Plan Crecimiento": {"precio": 170.0, "items": [
        "Todo lo del Plan Cuidado, más:",
        "1 ajuste mensual de optimización",
        "Revisión de formularios y CTA",
        "Ajuste de textos comerciales",
        "Métricas de conversión mensuales",
        "Reunión trimestral de estrategia"]},
    "Plan Escala": {"precio": 280.0, "items": [
        "Todo lo del Plan Crecimiento, más:",
        "2 ajustes mensuales de optimización",
        "Integración y revisión técnica de CRM",
        "Automatizaciones de seguimiento",
        "Revisión profunda de embudos",
        "Reporte avanzado de rendimiento"]},
    "Plan Liderazgo": {"precio": 450.0, "items": [
        "Todo lo del Plan Escala, más:",
        "Soporte técnico prioritario (SLA 4 hs)",
        "4 ajustes mensuales de optimización",
        "Reunión estratégica quincenal",
        "Gestión de campañas de automatización",
        "Mantenimiento de integraciones ERP complejas"]},
}


def mantenimiento_info(plan):
    return MANTENIMIENTO_INFO.get((plan or "Plan Cuidado").strip(),
                                  MANTENIMIENTO_INFO["Plan Cuidado"])


# ──────────────────────────────────────────────────────────────────────────────
# 3) NÚMERO A LETRAS (es) — para "monto en letras". Cubre enteros hasta millones.
# ──────────────────────────────────────────────────────────────────────────────
_UNID = ["", "uno", "dos", "tres", "cuatro", "cinco", "seis", "siete", "ocho", "nueve",
         "diez", "once", "doce", "trece", "catorce", "quince", "dieciséis",
         "diecisiete", "dieciocho", "diecinueve", "veinte", "veintiuno", "veintidós",
         "veintitrés", "veinticuatro", "veinticinco", "veintiséis", "veintisiete",
         "veintiocho", "veintinueve"]
_DECENAS = ["", "", "", "treinta", "cuarenta", "cincuenta", "sesenta", "setenta",
            "ochenta", "noventa"]
_CENTENAS = ["", "ciento", "doscientos", "trescientos", "cuatrocientos", "quinientos",
             "seiscientos", "setecientos", "ochocientos", "novecientos"]


def _centenas_a_letras(n: int) -> str:
    if n == 0:
        return ""
    if n == 100:
        return "cien"
    c, r = divmod(n, 100)
    out = _CENTENAS[c] if c else ""
    if r:
        if r < 30:
            dec = _UNID[r]
        else:
            d, u = divmod(r, 10)
            dec = _DECENAS[d] + (f" y {_UNID[u]}" if u else "")
        out = (out + " " + dec).strip()
    return out.strip()


def numero_a_letras(n: int) -> str:
    """Entero positivo → palabras en español. Fallback seguro a str(n)."""
    try:
        n = int(round(float(n)))
    except Exception:
        return str(n)
    if n == 0:
        return "cero"
    partes = []
    millones, resto = divmod(n, 1_000_000)
    miles, cent = divmod(resto, 1000)
    if millones:
        partes.append("un millón" if millones == 1
                       else f"{numero_a_letras(millones)} millones")
    if miles:
        partes.append("mil" if miles == 1 else f"{_centenas_a_letras(miles)} mil")
    if cent:
        partes.append(_centenas_a_letras(cent))
    return " ".join(p for p in partes if p).strip()


# ──────────────────────────────────────────────────────────────────────────────
# 4) MODELO DE DATOS DEL CONTRATO
# ──────────────────────────────────────────────────────────────────────────────
@dataclass
class DatosContrato:
    # Cliente (lo que el aliado completa en el mini-formulario)
    cliente_razon_social: str = ""
    cliente_cuit: str = ""
    cliente_domicilio: str = ""
    cliente_representante: str = ""
    cliente_cargo: str = ""
    cliente_email: str = ""
    cliente_pais: str = "Argentina"          # país del cliente (define la terminología fiscal)
    cliente_condicion_fiscal: str = ""       # situación tributaria (varía por país)
    cliente_dni: str = ""                    # documento del firmante (persona física)

    # Operación
    plan: str = "Plan Pro"
    precio_usd: Optional[float] = None          # si None, se toma de PLAN_INFO
    moneda: str = "USD"                          # "USD" | "ARS"
    precio_ars: Optional[float] = None
    tipo_cambio: Optional[float] = None
    factura_tipo: str = "B"                      # "A" | "B"
    iva_incluido: bool = True
    forma_pago: str = "pago único, sin costo mensual obligatorio"
    anticipo_pct: int = 100                      # 100 = todo por adelantado
    link_pago: str = ""

    # Mantenimiento opcional (Anexo II) — solo aparece si incluir_mantenimiento=True
    incluir_mantenimiento: bool = False
    plan_mantenimiento: str = "Plan Cuidado"
    mantenimiento_precio: Optional[float] = None   # si None, se toma de MANTENIMIENTO_INFO
    # Mora/punitorios opcional — útil cuando hay anticipo + saldo (anticipo_pct < 100)
    incluir_mora: bool = False
    interes_mora_mensual: float = 3.0              # % mensual

    # Lugar y fecha de firma
    ciudad: str = "Santa Fe"
    fecha: Optional[datetime.date] = None        # default: hoy

    # Overrides opcionales
    plazo_dias: Optional[int] = None
    garantia: Optional[str] = None
    capacitacion: Optional[int] = None
    entregables: Optional[list[str]] = None

    # Metadatos
    aliado_nombre: str = ""
    venta_id: Optional[int] = None
    numero_contrato: str = ""                # se autogenera si queda vacío

    def __post_init__(self):
        self.plan = normalizar_plan(self.plan)
        info = PLAN_INFO.get(self.plan, PLAN_INFO["Plan Pro"])
        if self.precio_usd is None:
            self.precio_usd = info["precio"]
        if self.plazo_dias is None:
            self.plazo_dias = info["plazo_dias"]
        if self.garantia is None:
            self.garantia = info["garantia"]
        if self.capacitacion is None:
            self.capacitacion = info["capacitacion"]
        if self.entregables is None:
            self.entregables = list(info["entregables"])
        if self.fecha is None:
            self.fecha = datetime.date.today()
        if not self.numero_contrato:
            _anio = (self.fecha or datetime.date.today()).year
            if self.venta_id:
                self.numero_contrato = f"AV-{_anio}-{int(self.venta_id):04d}"
            else:
                import random as _r
                self.numero_contrato = f"AV-{_anio}-{_r.randint(1000, 9999)}"
        if self.mantenimiento_precio is None:
            self.mantenimiento_precio = mantenimiento_info(self.plan_mantenimiento)["precio"]


def datos_desde_venta(venta, extra: Optional[dict] = None) -> DatosContrato:
    """
    Construye DatosContrato a partir de un objeto ORM `Venta`
    (models.Venta: nombre_cliente, plan, valor_usd, aliado...) más los campos
    fiscales que el aliado completa en `extra`.
    """
    extra = extra or {}
    aliado = getattr(venta, "aliado", None)
    fecha = extra.get("fecha")
    if isinstance(fecha, str) and fecha:
        try:
            fecha = datetime.date.fromisoformat(fecha)
        except Exception:
            fecha = None

    return DatosContrato(
        cliente_razon_social = extra.get("cliente_razon_social") or getattr(venta, "nombre_cliente", "") or "",
        cliente_cuit         = extra.get("cliente_cuit", ""),
        cliente_domicilio    = extra.get("cliente_domicilio", ""),
        cliente_representante= extra.get("cliente_representante", ""),
        cliente_cargo        = extra.get("cliente_cargo", ""),
        cliente_email        = extra.get("cliente_email", ""),
        cliente_pais             = extra.get("cliente_pais", "Argentina"),
        cliente_condicion_fiscal = extra.get("cliente_condicion_fiscal", ""),
        cliente_dni              = extra.get("cliente_dni", ""),
        numero_contrato          = extra.get("numero_contrato", ""),
        incluir_mantenimiento    = bool(extra.get("incluir_mantenimiento", False)),
        plan_mantenimiento       = extra.get("plan_mantenimiento", "Plan Cuidado") or "Plan Cuidado",
        mantenimiento_precio     = extra.get("mantenimiento_precio"),
        incluir_mora             = bool(extra.get("incluir_mora", False)),
        interes_mora_mensual     = float(extra.get("interes_mora_mensual", 3.0) or 3.0),
        plan                 = getattr(venta, "plan", "Plan Pro") or "Plan Pro",
        precio_usd           = extra.get("precio_usd", getattr(venta, "valor_usd", None)),
        moneda               = extra.get("moneda", "USD"),
        precio_ars           = extra.get("precio_ars"),
        tipo_cambio          = extra.get("tipo_cambio"),
        factura_tipo         = extra.get("factura_tipo", "B"),
        iva_incluido         = bool(extra.get("iva_incluido", True)),
        forma_pago           = extra.get("forma_pago", "pago único, sin costo mensual obligatorio"),
        anticipo_pct         = int(extra.get("anticipo_pct", 100)),
        link_pago            = extra.get("link_pago", ""),
        ciudad               = extra.get("ciudad", "Santa Fe"),
        fecha                = fecha,
        plazo_dias           = extra.get("plazo_dias"),
        garantia             = extra.get("garantia"),
        capacitacion         = extra.get("capacitacion"),
        entregables          = extra.get("entregables"),
        aliado_nombre        = getattr(aliado, "nombre", "") if aliado else extra.get("aliado_nombre", ""),
        venta_id             = getattr(venta, "id", None),
    )


# ──────────────────────────────────────────────────────────────────────────────
# 5) HELPERS DE PRESENTACIÓN
# ──────────────────────────────────────────────────────────────────────────────
def _esc(s) -> str:
    s = "" if s is None else str(s)
    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def _campo(valor: str, etiqueta: str) -> str:
    """Muestra el valor o, si está vacío, un marcador a completar a mano (HTML)."""
    valor = (valor or "").strip()
    if valor:
        return f"<strong>{_esc(valor)}</strong>"
    return f'<span class="pendiente">[completar: {_esc(etiqueta)}]</span>'


def _es_placeholder(v) -> bool:
    """True si el dato de Avanza no fue cargado (vacío o placeholder '[...]')."""
    v = (v or "").strip()
    return (not v) or v.startswith("[")


def _avanza_campo(valor, etiqueta) -> str:
    """Igual que _campo pero para datos del PRESTADOR: si falta la env var,
    se marca en rojo para que NUNCA salga un contrato con '[CUIT de Avanza]'."""
    if _es_placeholder(valor):
        return f'<span class="pendiente">[completar config: {_esc(etiqueta)}]</span>'
    return _esc(valor)


def _avanza_txt(valor, etiqueta) -> str:
    """Versión texto plano (para DOCX)."""
    if _es_placeholder(valor):
        return f"«COMPLETAR CONFIG: {etiqueta}»"
    return str(valor)


def _inline_html(text: str) -> str:
    """Escapa y convierte **negrita** -> <strong>negrita</strong>."""
    out, partes = [], _esc(text).split("**")
    for i, frag in enumerate(partes):
        out.append(f"<strong>{frag}</strong>" if i % 2 == 1 else frag)
    return "".join(out)


def _precio_str(d) -> str:
    usd = f"USD {d.precio_usd:,.0f}".replace(",", ".")
    letras = numero_a_letras(d.precio_usd)
    base = f"{usd} (dólares estadounidenses {letras})"
    if d.moneda == "ARS" and d.precio_ars:
        ars = f"ARS {d.precio_ars:,.0f}".replace(",", ".")
        tc = ""
        if d.tipo_cambio:
            tc = f" — tipo de cambio de referencia: {d.tipo_cambio:,.2f}"
            tc = tc.replace(",", "@").replace(".", ",").replace("@", ".")
        base += f", pagadero en pesos por {ars}{tc}"
    return base


def _logo_html() -> str:
    path = AVANZA["logo_path"]
    if path and (path.startswith("data:") or os.path.exists(path)):
        src = path if path.startswith("data:") else f"file://{os.path.abspath(path)}"
        return f'<img class="logo-img" src="{src}" alt="Avanza Digital">'
    return ('<div class="wordmark">» Avanza <span>Digital</span></div>'
            '<div class="tagline">Sistemas comerciales automatizados B2B</div>')


# ──────────────────────────────────────────────────────────────────────────────
# CONTENIDO COMPARTIDO — una sola fuente de verdad para PDF y Word
# ──────────────────────────────────────────────────────────────────────────────
def _derivados(d) -> dict:
    info = PLAN_INFO.get(d.plan, PLAN_INFO["Plan Pro"])
    iva = ("Los importes consignados incluyen el IVA correspondiente." if d.iva_incluido
           else "A los importes consignados deberá adicionarse el IVA que corresponda según la normativa vigente.")
    cambio = ("" if d.moneda == "ARS" else
              " En caso de abonarse en pesos argentinos, la conversión se realizará al tipo de cambio vendedor del "
              "Banco de la Nación Argentina vigente al día del efectivo pago.")
    sla = (" Para el plan contratado se establece un Acuerdo de Nivel de Servicio (SLA) de respuesta de hasta 24 "
           "(veinticuatro) horas hábiles ante incidencias críticas reportadas durante el período de garantía."
           if info.get("sla") else "")
    cap = d.capacitacion or 0
    cap_txt = (f"AVANZA brindará {cap} ({numero_a_letras(cap)}) sesión(es) de capacitación al equipo de EL CLIENTE "
               "sobre el uso y la operación del sistema entregado." if cap > 0 else
               "El plan contratado no incluye sesiones de capacitación presenciales; AVANZA entregará la documentación "
               "de uso correspondiente. Sesiones adicionales podrán contratarse por separado.")
    anticipo = ("El precio se abonará en un 100% (cien por ciento) por adelantado, contra emisión de la factura, dando "
                "inicio la implementación una vez acreditado el pago." if d.anticipo_pct >= 100 else
                f"El precio se abonará en un {d.anticipo_pct}% en concepto de anticipo a la firma y el "
                f"{100 - d.anticipo_pct}% restante contra entrega y aceptación del sistema. La implementación dará "
                "comienzo una vez acreditado el anticipo.")
    link = (f" a través del enlace de pago provisto por el asesor de AVANZA ({d.link_pago})" if d.link_pago else
            " a través del enlace de pago provisto por el asesor de AVANZA, o por el medio que AVANZA indique")
    pais_norm = _strip_accents(d.cliente_pais or "Argentina").strip().lower()
    if pais_norm in ("argentina", "", "ar"):
        comprobante = f"la factura tipo **{d.factura_tipo}**"
    else:
        comprobante = "el comprobante fiscal que corresponda según la normativa vigente"
    return {"iva": iva, "cambio": cambio, "sla": sla, "cap": cap_txt, "anticipo": anticipo, "link": link, "comprobante": comprobante}


def _clausulas(d) -> list:
    """Lista de (título, [párrafos]). Los párrafos pueden traer **negrita**."""
    x = _derivados(d)
    precio = _precio_str(d)
    return [
        ("CLÁUSULA PRIMERA — OBJETO", [
            f"AVANZA se obliga a diseñar, desarrollar e implementar para EL CLIENTE un sistema comercial digital "
            f"correspondiente al **{d.plan}**, con el alcance y los entregables que se detallan en la Cláusula Segunda "
            f"y en el Anexo I, que forma parte integrante de este Contrato."]),
        ("CLÁUSULA SEGUNDA — ALCANCE Y ENTREGABLES", [
            "El servicio comprende los entregables correspondientes al plan contratado, detallados en el **Anexo I**. "
            "Cualquier desarrollo, funcionalidad o requerimiento no incluido expresamente en dicho anexo será "
            "considerado fuera de alcance y, de ser solicitado, presupuestado por separado."]),
        ("CLÁUSULA TERCERA — PLAZO DE IMPLEMENTACIÓN", [
            f"AVANZA entregará el sistema en un plazo estimado de **{d.plazo_dias} días hábiles**, contados a partir de "
            "(i) la acreditación del pago conforme a la Cláusula Cuarta y (ii) la entrega por parte de EL CLIENTE de la "
            "totalidad de los contenidos, accesos y aprobaciones necesarios. Las demoras imputables a EL CLIENTE en la "
            "provisión de dichos elementos extenderán el plazo en igual medida."]),
        ("CLÁUSULA CUARTA — PRECIO Y FORMA DE PAGO", [
            f"El precio total del servicio es de **{precio}**, en concepto de {d.forma_pago}. {x['iva']}{x['cambio']}",
            f"{x['anticipo']} El pago se efectuará{x['link']}. EL CLIENTE recibirá {x['comprobante']} "
            "correspondiente.",
            "Nota: los planes de mantenimiento y evolución posteriores son opcionales y se contratan por separado, sin "
            "permanencia mínima; no son requisito para el funcionamiento del sistema."]),
        ("CLÁUSULA QUINTA — ACEPTACIÓN DE ENTREGABLES", [
            "Entregado el sistema, EL CLIENTE dispondrá de un plazo de 7 (siete) días corridos para revisarlo y "
            "notificar por medio fehaciente las observaciones que correspondan al alcance del Anexo I. El plan incluye "
            "hasta 2 (dos) rondas de ajustes sobre dichas observaciones. Transcurrido el plazo sin observaciones, o "
            "puesto el sistema en uso productivo por EL CLIENTE, los entregables se tendrán por aceptados de "
            "conformidad, iniciándose el cómputo del período de garantía."]),
        ("CLÁUSULA SEXTA — PROPIEDAD DEL CÓDIGO Y LOS ENTREGABLES", [
            "Una vez abonado el precio total, los derechos patrimoniales sobre el **desarrollo a medida** realizado "
            "específicamente para EL CLIENTE (código fuente y entregables) quedan en propiedad del 100% de EL CLIENTE, "
            "sin licencias de uso, alquileres ni pagos recurrentes obligatorios a favor de AVANZA. Quedan exceptuados "
            "los componentes, librerías, frameworks y software de terceros (incluyendo soluciones de código abierto), "
            "que se entregan y continúan rigiéndose bajo sus respectivas licencias de origen, sin que su titularidad "
            "pueda ser cedida. AVANZA conserva la propiedad de su conocimiento técnico (know-how) y de los componentes "
            "genéricos y reutilizables de su autoría. AVANZA entregará el código y los accesos correspondientes al "
            "finalizar la implementación."]),
        ("CLÁUSULA SÉPTIMA — HOSTING, DOMINIO, SSL Y CORREO", [
            "El precio del plan **no incluye** el alojamiento (hosting), el registro o renovación del dominio, el "
            "certificado de seguridad (SSL) ni las casillas de correo corporativo. EL CLIENTE podrá (i) alojar el "
            "sistema en su propia infraestructura, o (ii) contratar el plan de mantenimiento e infraestructura de "
            "AVANZA (Plan Cuidado), de carácter mensual y opcional. El dominio se registrará a nombre de EL CLIENTE. "
            "Hasta tanto EL CLIENTE provea su entorno de alojamiento o contrate dicho plan, AVANZA no garantiza la "
            "disponibilidad en línea del sistema."]),
        ("CLÁUSULA OCTAVA — SERVICIOS Y SOFTWARE DE TERCEROS", [
            "El funcionamiento de determinadas funcionalidades puede depender de servicios, plataformas o software de "
            "terceros (por ejemplo: pasarelas de pago, WhatsApp Business API, servicios de inteligencia artificial, "
            "Google, y sistemas de gestión o ERP como Tango o SAP). La contratación, las licencias, las cuentas y los "
            "costos —iniciales o recurrentes— de dichos terceros son a cargo exclusivo de EL CLIENTE. AVANZA no "
            "responde por cambios de precio, condiciones, disponibilidad o discontinuación de tales servicios de "
            "terceros."]),
        ("CLÁUSULA NOVENA — OBLIGACIONES DE LAS PARTES", [
            "**De AVANZA:** ejecutar el servicio conforme al alcance del Anexo I y en el plazo pactado; entregar el "
            "código fuente y la documentación de acceso al finalizar; y brindar la capacitación y el soporte de "
            "garantía previstos en este Contrato.",
            "**De EL CLIENTE:** abonar el precio en la forma y plazo establecidos; proveer en tiempo los contenidos, "
            "accesos, marca y aprobaciones necesarios; y designar un responsable de contacto para la coordinación del "
            "proyecto."]),
        ("CLÁUSULA DÉCIMA — CAPACITACIÓN", [x["cap"]]),
        ("CLÁUSULA UNDÉCIMA — GARANTÍA Y SOPORTE", [
            f"AVANZA garantiza el correcto funcionamiento del sistema entregado por un período de **{d.garantia}** desde "
            f"la aceptación, corrigiendo sin cargo los defectos atribuibles al desarrollo.{x['sla']} Quedan excluidos "
            "de la garantía los daños derivados de modificaciones realizadas por terceros, mal uso, o causas ajenas a "
            "AVANZA. Finalizado el período de garantía, las correcciones, el mantenimiento y la evolución podrán "
            "contratarse mediante bolsas de horas o planes de mantenimiento mensual, a la tarifa horaria vigente de "
            "AVANZA."]),
        ("CLÁUSULA DUODÉCIMA — PROTECCIÓN DE DATOS PERSONALES", [
            "Las Partes se obligan a tratar los datos personales conforme a la Ley Nacional N.º 25.326 de Protección de "
            "Datos Personales y normativa concordante. Respecto de los datos de los contactos y clientes finales "
            "captados a través del sistema, EL CLIENTE reviste el carácter de responsable de la base de datos, y AVANZA "
            "actúa como encargada del tratamiento únicamente durante la implementación y en la medida necesaria para "
            "prestar el servicio, sin utilizarlos para fines propios. Finalizado el Contrato, AVANZA pondrá tales datos "
            "a disposición de EL CLIENTE y procederá a su eliminación de sus entornos de trabajo, salvo obligación "
            "legal de conservación."]),
        ("CLÁUSULA DECIMOTERCERA — CONFIDENCIALIDAD", [
            "Las Partes se obligan a mantener la confidencialidad de toda información a la que accedan con motivo de "
            "este Contrato, y a no divulgarla ni utilizarla para fines distintos a su ejecución, obligación que "
            "subsistirá tras la finalización del Contrato."]),
        ("CLÁUSULA DECIMOCUARTA — LIMITACIÓN DE RESPONSABILIDAD", [
            "La responsabilidad total de AVANZA frente a EL CLIENTE, por cualquier concepto derivado de este Contrato, "
            "se limita al monto efectivamente abonado por EL CLIENTE. AVANZA no responderá por lucro cesante ni por "
            "daños indirectos. Las garantías de disponibilidad (uptime) o niveles de servicio publicadas con fines "
            "comerciales corresponden exclusivamente a los planes de mantenimiento mensual y no resultan exigibles bajo "
            "la modalidad de pago único aquí contratada."]),
        ("CLÁUSULA DECIMOQUINTA — INDEMNIDAD POR CONTENIDOS DE EL CLIENTE", [
            "EL CLIENTE declara que los contenidos, marcas, logotipos, textos e imágenes que provea para el desarrollo "
            "son de su titularidad o cuenta con licencia suficiente para su uso, y mantendrá indemne a AVANZA frente a "
            "cualquier reclamo de terceros derivado de dichos materiales."]),
        ("CLÁUSULA DECIMOSEXTA — FUERZA MAYOR", [
            "Ninguna de las Partes será responsable por el incumplimiento o demora derivados de caso fortuito o fuerza "
            "mayor. La Parte afectada lo notificará a la otra a la brevedad, suspendiéndose los plazos mientras dure el "
            "impedimento."]),
        ("CLÁUSULA DECIMOSÉPTIMA — CESIÓN", [
            "Ninguna de las Partes podrá ceder su posición contractual ni los derechos y obligaciones emergentes de "
            "este Contrato sin el consentimiento previo y por escrito de la otra Parte."]),
        ("CLÁUSULA DECIMOCTAVA — VIGENCIA Y RESCISIÓN", [
            "El Contrato rige desde su firma hasta el cumplimiento total de las obligaciones de las Partes y la "
            "finalización del período de garantía. Cualquiera de las Partes podrá rescindirlo ante el incumplimiento "
            "grave de la otra, previa intimación por medio fehaciente y un plazo de 10 (diez) días para subsanar. En "
            "caso de rescisión por causa imputable a EL CLIENTE luego de iniciada la implementación, AVANZA podrá "
            "retener los montos correspondientes a las tareas ya ejecutadas."]),
        ("CLÁUSULA DECIMONOVENA — MODIFICACIONES", [
            "Toda modificación al presente Contrato deberá constar por escrito y ser firmada por ambas Partes para "
            "tener validez."]),
        ("CLÁUSULA VIGÉSIMA — LEY APLICABLE Y JURISDICCIÓN", [
            "El presente Contrato se rige por las leyes de la República Argentina. Para toda controversia derivada del "
            f"mismo, las Partes se someten a la jurisdicción de los tribunales ordinarios de {AVANZA['jurisdiccion']}, "
            "renunciando a cualquier otro fuero o jurisdicción que pudiera corresponder."]),
        ("CLÁUSULA VIGÉSIMA PRIMERA — NOTIFICACIONES", [
            "Las Partes acuerdan que todas las notificaciones vinculadas a este Contrato se tendrán por válidas "
            "si se cursan a las direcciones de correo electrónico que declaran a tal efecto (AVANZA: "
            f"{AVANZA['email']}" + (f"; EL CLIENTE: {d.cliente_email}" if d.cliente_email else "") + ")."]),
        ("CLÁUSULA VIGÉSIMA SEGUNDA — INSTRUMENTACIÓN Y FIRMA", [
            "El presente Contrato podrá instrumentarse y celebrarse válidamente por medios electrónicos. Las "
            "Partes acuerdan que su firma mediante firma electrónica o firma digital, así como la aceptación por "
            "intercambio de correos electrónicos o la suscripción de ejemplares en formato PDF remitidos entre las "
            "Partes, tendrá plena validez y eficacia jurídica y se tendrá por expresión auténtica del "
            "consentimiento, conforme a la Ley N.º 25.506 de Firma Digital y a los artículos 286 y 288 del Código "
            "Civil y Comercial de la Nación. Las Partes renuncian a objetar la validez o fuerza probatoria del "
            "Contrato por el solo hecho de haberse celebrado o suscripto por medios electrónicos."]),
    ] + ([
        ("CLÁUSULA VIGÉSIMA TERCERA — MORA E INTERESES", [
            f"La falta de pago en término de cualquier suma adeudada —en particular, el saldo posterior al "
            f"anticipo previsto en la Cláusula Cuarta— devengará de pleno derecho, sin necesidad de interpelación "
            f"previa, un interés moratorio del {d.interes_mora_mensual:g}% (por ciento) mensual sobre el importe "
            "impago, desde su vencimiento y hasta su efectiva cancelación. La mora faculta a AVANZA a suspender "
            "la implementación hasta la regularización del pago, sin que ello genere responsabilidad a su cargo."])]
        if d.incluir_mora else [])


# ──────────────────────────────────────────────────────────────────────────────
# CSS (impresión / xhtml2pdf)
# ──────────────────────────────────────────────────────────────────────────────
CSS = """
@page {
  size: A4;
  margin: 22mm 20mm 28mm 20mm;
  @frame footer {
    -pdf-frame-content: pdf_footer;
    bottom: 8pt;
    left: 20mm;
    right: 20mm;
    height: 16pt;
  }
}
* { box-sizing: border-box; }
body { font-family: "DejaVu Sans", "Inter", Arial, sans-serif; font-size: 9.6pt; line-height: 1.5; color: #1f2733; margin: 0; }
.header { border-bottom: 2px solid #1463ff; padding-bottom: 10px; margin-bottom: 16px; }
.logo-img { height: 34px; }
.wordmark { font-size: 18pt; font-weight: 800; color: #0b1b3a; letter-spacing: -.5px; }
.wordmark span { color: #1463ff; }
.tagline { font-size: 8pt; color: #6b7280; letter-spacing: 2px; text-transform: uppercase; margin-top: 2px; }
h1 { font-size: 13pt; text-align: center; color: #0b1b3a; margin: 10px 0 2px; letter-spacing: .3px; }
.subtitulo { text-align: center; font-size: 8.5pt; color: #6b7280; margin: 0 0 16px; }
.intro { text-align: justify; margin-bottom: 12px; }
.partes { background: #f5f8ff; border: 1px solid #dbe6ff; border-radius: 6px; padding: 10px 12px; margin-bottom: 14px; }
.partes p { margin: 4px 0; }
h2 { font-size: 9.6pt; color: #0b1b3a; margin: 14px 0 4px; border-left: 3px solid #1463ff; padding-left: 8px; }
p { margin: 5px 0; text-align: justify; }
ul { margin: 4px 0 8px 0; padding-left: 18px; }
li { margin: 2px 0; }
.pendiente { color: #c2410c; background: #fff4ed; padding: 0 3px; border-radius: 3px; font-weight: 600; }
.firma-tabla { width: 100%; border-collapse: collapse; margin-top: 26px; }
.firma-tabla td { width: 50%; border: 1px solid #cbd5e1; vertical-align: top; padding: 14px 12px; font-size: 8.6pt; }
.firma-tabla .rol { font-weight: 700; color: #0b1b3a; margin-bottom: 26px; }
.firma-linea { border-top: 1px solid #94a3b8; margin-top: 22px; padding-top: 3px; }
.anexo { page-break-before: always; }
.anexo h1 { text-align: left; font-size: 12pt; border-bottom: 1px solid #dbe6ff; padding-bottom: 6px; }
.anexo-meta { background: #f5f8ff; border: 1px solid #dbe6ff; border-radius: 6px; padding: 8px 12px; margin: 8px 0 12px; font-size: 9pt; }
.disclaimer { margin-top: 16px; font-size: 7.6pt; color: #94a3b8; font-style: italic; text-align: center; }
.nrocontrato { text-align: center; font-size: 8.5pt; color: #1463ff; font-weight: 700; margin: 0 0 12px; letter-spacing: .5px; }
"""


# ──────────────────────────────────────────────────────────────────────────────
# RENDER HTML (para PDF vía WeasyPrint)
# ──────────────────────────────────────────────────────────────────────────────
def render_html(d) -> str:
    fecha = d.fecha
    dia, mes, anio = fecha.day, MESES_ES[fecha.month], fecha.year
    cargo = f" — {_esc(d.cliente_cargo)}" if d.cliente_cargo else ""
    ident = ident_fiscal(d.cliente_pais)
    cond_html = f" ({_esc(d.cliente_condicion_fiscal)})" if d.cliente_condicion_fiscal else ""
    pais_html = f", {_esc(d.cliente_pais)}" if (d.cliente_pais and _strip_accents(d.cliente_pais).strip().lower() not in ("argentina", "")) else ""
    dni_html = f" ({_esc(ident['persona'])}: {_campo(d.cliente_dni, ident['persona'] + ' del firmante')})"
    anexo2_html = ""
    if d.incluir_mantenimiento:
        _mi = mantenimiento_info(d.plan_mantenimiento)
        _pm = d.mantenimiento_precio if d.mantenimiento_precio is not None else _mi["precio"]
        _items2 = "".join(f"<li>{_esc(it)}</li>" for it in _mi["items"])
        anexo2_html = (
            '<div class="anexo"><h1>ANEXO II — PLAN DE MANTENIMIENTO (OPCIONAL)</h1>'
            f'<div class="anexo-meta"><strong>Plan:</strong> {_esc(d.plan_mantenimiento)} &nbsp;·&nbsp;'
            f'<strong>Abono mensual:</strong> USD {_pm:,.0f} / mes &nbsp;·&nbsp;'
            '<strong>Permanencia:</strong> sin permanencia mínima</div>'
            '<p>El presente plan de mantenimiento es <strong>opcional</strong> y se contrata por separado del '
            'servicio de la Cláusula Cuarta. No es requisito para el funcionamiento del sistema entregado. '
            f'Incluye:</p><ul>{_items2}</ul>'
            '<p style="font-size:8pt;color:#6b7280;">El abono podrá actualizarse periódicamente. EL CLIENTE '
            'podrá darlo de baja en cualquier momento sin penalidad, con un preaviso de 30 (treinta) días.</p></div>')
    clausulas_html = "".join(
        f"<h2>{_esc(t)}</h2>" + "".join(f"<p>{_inline_html(p)}</p>" for p in cuerpo)
        for t, cuerpo in _clausulas(d)
    )
    entregables_html = "".join(f"<li>{_esc(e)}</li>" for e in d.entregables)

    return f"""<!DOCTYPE html>
<html lang="es"><head><meta charset="utf-8"><style>{CSS}</style></head>
<body>
<div class="header">{_logo_html()}</div>
<h1>CONTRATO DE PRESTACIÓN DE SERVICIOS</h1>
<p class="subtitulo">Implementación de sistema comercial digital</p>
<p class="nrocontrato">Contrato N.º {_esc(d.numero_contrato)}</p>
<p class="intro">En la ciudad de {_campo(d.ciudad, "ciudad")}, a los <strong>{dia}</strong> días del mes de
<strong>{mes}</strong> de <strong>{anio}</strong>, entre las partes que a continuación se identifican, se celebra el
presente Contrato de Prestación de Servicios (en adelante, el «Contrato»):</p>
<div class="partes">
  <p><strong>EL PRESTADOR:</strong> {_esc(AVANZA['razon_social'])}, CUIT {_avanza_campo(AVANZA['cuit'], 'CUIT de Avanza · env AVANZA_CUIT')}, representada en este acto por {_avanza_campo(AVANZA['representante'], 'representante · env AVANZA_REPRESENTANTE')}
     ({_esc(AVANZA['cargo'])}) (en adelante, «AVANZA»).</p>
  <p><strong>EL CLIENTE:</strong> {_campo(d.cliente_razon_social, "razón social del cliente")},
     {_esc(ident['empresa'])} {_campo(d.cliente_cuit, ident['empresa'] + " del cliente")}{cond_html}{pais_html},
     representada en este acto por {_campo(d.cliente_representante, "representante")}{dni_html}{cargo} (en adelante, «EL CLIENTE»).</p>
</div>
<p>AVANZA y EL CLIENTE se denominarán conjuntamente «las Partes». Las Partes acuerdan celebrar el presente Contrato
conforme a las siguientes cláusulas:</p>
{clausulas_html}
<p style="margin-top:12px;">En prueba de conformidad, las Partes suscriben el presente —sea en dos ejemplares de un
mismo tenor y a un solo efecto, sea por medios electrónicos conforme la Cláusula de Instrumentación y Firma— en el
lugar y fecha indicados.</p>
<table class="firma-tabla"><tr>
  <td><div class="rol">P/ AVANZA DIGITAL</div>
    <div class="firma-linea">Aclaración: {_avanza_campo(AVANZA['representante'], 'representante · env AVANZA_REPRESENTANTE')}</div>
    <div>Cargo: {_esc(AVANZA['cargo'])}</div><div>CUIT: {_avanza_campo(AVANZA['cuit'], 'CUIT de Avanza · env AVANZA_CUIT')}</div>
    <div>Fecha: ____ / ____ / ______</div></td>
  <td><div class="rol">P/ EL CLIENTE</div>
    <div class="firma-linea">Aclaración: {_esc(d.cliente_representante) or '________________________'}</div>
    <div>Cargo: {_esc(d.cliente_cargo) or '________________________'}</div>
    <div>{_esc(ident['empresa'])}: {_esc(d.cliente_cuit) or '________________________'}</div>
    <div>{_esc(ident['persona'])}: {_esc(d.cliente_dni) or '________________________'}</div>
    <div>Fecha: ____ / ____ / ______</div></td>
</tr></table>
<div class="anexo">
  <h1>ANEXO I — DETALLE DEL PLAN CONTRATADO</h1>
  <div class="anexo-meta"><strong>Plan:</strong> {_esc(d.plan)} &nbsp;·&nbsp;
    <strong>Precio:</strong> USD {d.precio_usd:,.0f} &nbsp;·&nbsp;
    <strong>Plazo:</strong> {d.plazo_dias} días hábiles &nbsp;·&nbsp;
    <strong>Garantía:</strong> {_esc(d.garantia)}</div>
  <p><strong>Entregables incluidos:</strong></p>
  <ul>{entregables_html}</ul>
  <p class="disclaimer">Documento generado automáticamente por Avanza Digital. Modelo de base; ante dudas, consulte
  asesoramiento legal.</p>
</div>
{anexo2_html}
<!-- Footer fijo para xhtml2pdf -->
<div id="pdf_footer">
  <table style="width:100%;border-collapse:collapse;">
    <tr>
      <td style="font-size:7.5pt;color:#9aa3b2;letter-spacing:.3px;">
        Avanza Digital &middot; avanzadigital.digital &middot; Santa Fe, Argentina
      </td>
      <td style="font-size:7.5pt;color:#9aa3b2;text-align:right;">
        P&aacute;gina <pdf:pagenumber> de <pdf:pagecount>
      </td>
    </tr>
  </table>
</div>
</body></html>"""


# ──────────────────────────────────────────────────────────────────────────────
# RENDER DOCX (Word editable, vía python-docx)
# ──────────────────────────────────────────────────────────────────────────────
def _docx_runs(paragraph, text):
    """Agrega runs a un párrafo interpretando **negrita**."""
    for i, frag in enumerate(text.split("**")):
        if frag == "":
            continue
        run = paragraph.add_run(frag)
        run.bold = (i % 2 == 1)


def _docx_campo(paragraph, valor, etiqueta):
    from docx.shared import RGBColor
    valor = (valor or "").strip()
    if valor:
        paragraph.add_run(valor).bold = True
    else:
        r = paragraph.add_run(f"[completar: {etiqueta}]")
        r.bold = True
        r.font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)


def _docx_page_field(paragraph):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    run = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = 'PAGE'
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    run._r.append(b); run._r.append(i); run._r.append(e)


def render_contrato_docx(d) -> bytes:
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    fecha = d.fecha
    dia, mes, anio = fecha.day, MESES_ES[fecha.month], fecha.year
    AZUL = RGBColor(0x14, 0x63, 0xFF)
    NAVY = RGBColor(0x0B, 0x1B, 0x3A)

    doc = Document()
    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(10.5)

    # Encabezado / wordmark
    h = doc.add_paragraph()
    r1 = h.add_run("» Avanza "); r1.bold = True; r1.font.size = Pt(20); r1.font.color.rgb = NAVY
    r2 = h.add_run("Digital");   r2.bold = True; r2.font.size = Pt(20); r2.font.color.rgb = AZUL
    tg = doc.add_paragraph(); tr = tg.add_run("SISTEMAS COMERCIALES AUTOMATIZADOS B2B")
    tr.font.size = Pt(8); tr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    titulo = doc.add_paragraph(); titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tt = titulo.add_run("CONTRATO DE PRESTACIÓN DE SERVICIOS"); tt.bold = True; tt.font.size = Pt(14); tt.font.color.rgb = NAVY
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Implementación de sistema comercial digital"); sr.font.size = Pt(9); sr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    nro = doc.add_paragraph(); nro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nrr = nro.add_run(f"Contrato N.º {d.numero_contrato}"); nrr.bold = True; nrr.font.size = Pt(8.5); nrr.font.color.rgb = AZUL

    intro = doc.add_paragraph()
    intro.add_run("En la ciudad de ")
    _docx_campo(intro, d.ciudad, "ciudad")
    intro.add_run(f", a los {dia} días del mes de {mes} de {anio}, entre las partes que a continuación se identifican, "
                  "se celebra el presente Contrato de Prestación de Servicios (en adelante, el «Contrato»):")

    _av_cuit = _avanza_txt(AVANZA["cuit"], "CUIT de Avanza · env AVANZA_CUIT")
    _av_rep  = _avanza_txt(AVANZA["representante"], "representante · env AVANZA_REPRESENTANTE")
    p1 = doc.add_paragraph(); p1.add_run("EL PRESTADOR: ").bold = True
    p1.add_run(f"{AVANZA['razon_social']}, CUIT {_av_cuit}, representada "
               f"en este acto por {_av_rep} ({AVANZA['cargo']}) (en adelante, «AVANZA»).")
    ident = ident_fiscal(d.cliente_pais)
    p2 = doc.add_paragraph(); p2.add_run("EL CLIENTE: ").bold = True
    _docx_campo(p2, d.cliente_razon_social, "razón social del cliente")
    p2.add_run(f", {ident['empresa']} ")
    _docx_campo(p2, d.cliente_cuit, ident['empresa'] + " del cliente")
    if d.cliente_condicion_fiscal:
        p2.add_run(f" ({d.cliente_condicion_fiscal})")
    if d.cliente_pais and _strip_accents(d.cliente_pais).strip().lower() not in ("argentina", ""):
        p2.add_run(f", {d.cliente_pais}")
    p2.add_run(", representada en este acto por ")
    _docx_campo(p2, d.cliente_representante, "representante")
    p2.add_run(f" ({ident['persona']}: ")
    _docx_campo(p2, d.cliente_dni, ident['persona'] + " del firmante")
    p2.add_run(")")
    if d.cliente_cargo:
        p2.add_run(f" — {d.cliente_cargo}")
    p2.add_run(" (en adelante, «EL CLIENTE»).")

    doc.add_paragraph("AVANZA y EL CLIENTE se denominarán conjuntamente «las Partes». Las Partes acuerdan celebrar el "
                      "presente Contrato conforme a las siguientes cláusulas:")

    for titulo_cl, cuerpo in _clausulas(d):
        hp = doc.add_paragraph(); hp.paragraph_format.space_before = Pt(8)
        hr = hp.add_run(titulo_cl); hr.bold = True; hr.font.color.rgb = NAVY; hr.font.size = Pt(10.5)
        for parr in cuerpo:
            pp = doc.add_paragraph(); _docx_runs(pp, parr)

    doc.add_paragraph("En prueba de conformidad, las Partes suscriben el presente —sea en dos ejemplares de un mismo "
                      "tenor y a un solo efecto, sea por medios electrónicos conforme la Cláusula de Instrumentación y "
                      "Firma— en el lugar y fecha indicados.").paragraph_format.space_before = Pt(10)

    # Tabla de firmas
    tabla = doc.add_table(rows=1, cols=2); tabla.style = "Table Grid"
    izq, der = tabla.rows[0].cells
    izq.paragraphs[0].add_run("P/ AVANZA DIGITAL").bold = True
    for t in [f"\nAclaración: {_av_rep}", f"Cargo: {AVANZA['cargo']}",
              f"CUIT: {_av_cuit}", "Fecha: ____ / ____ / ______"]:
        izq.add_paragraph(t)
    der.paragraphs[0].add_run("P/ EL CLIENTE").bold = True
    for t in [f"\nAclaración: {d.cliente_representante or '________________________'}",
              f"Cargo: {d.cliente_cargo or '________________________'}",
              f"{ident['persona']}: {d.cliente_dni or '________________________'}",
              f"{ident['empresa']}: {d.cliente_cuit or '________________________'}", "Fecha: ____ / ____ / ______"]:
        der.add_paragraph(t)

    # Anexo I (página nueva)
    doc.add_page_break()
    ax = doc.add_paragraph(); axr = ax.add_run("ANEXO I — DETALLE DEL PLAN CONTRATADO")
    axr.bold = True; axr.font.size = Pt(12); axr.font.color.rgb = NAVY
    meta = doc.add_paragraph()
    meta.add_run("Plan: ").bold = True; meta.add_run(f"{d.plan}    ·    ")
    meta.add_run("Precio: ").bold = True; meta.add_run(f"USD {d.precio_usd:,.0f}    ·    ")
    meta.add_run("Plazo: ").bold = True; meta.add_run(f"{d.plazo_dias} días hábiles    ·    ")
    meta.add_run("Garantía: ").bold = True; meta.add_run(d.garantia)
    doc.add_paragraph().add_run("Entregables incluidos:").bold = True
    for e in d.entregables:
        doc.add_paragraph(e, style="List Bullet")
    dis = doc.add_paragraph(); dr = dis.add_run(
        "Documento generado automáticamente por Avanza Digital. Modelo de base; ante dudas, consulte asesoramiento legal.")
    dr.italic = True; dr.font.size = Pt(8); dr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    if d.incluir_mantenimiento:
        _mi = mantenimiento_info(d.plan_mantenimiento)
        _pm = d.mantenimiento_precio if d.mantenimiento_precio is not None else _mi["precio"]
        doc.add_page_break()
        a2 = doc.add_paragraph(); a2r = a2.add_run("ANEXO II — PLAN DE MANTENIMIENTO (OPCIONAL)")
        a2r.bold = True; a2r.font.size = Pt(12); a2r.font.color.rgb = NAVY
        m2 = doc.add_paragraph()
        m2.add_run("Plan: ").bold = True; m2.add_run(f"{d.plan_mantenimiento}    ·    ")
        m2.add_run("Abono mensual: ").bold = True; m2.add_run(f"USD {_pm:,.0f} / mes    ·    ")
        m2.add_run("Permanencia: ").bold = True; m2.add_run("sin permanencia mínima")
        doc.add_paragraph("Este plan de mantenimiento es opcional y se contrata por separado del servicio de la "
                          "Cláusula Cuarta; no es requisito para el funcionamiento del sistema. Incluye:")
        for _it in _mi["items"]:
            doc.add_paragraph(_it, style="List Bullet")
        n2 = doc.add_paragraph(); n2r = n2.add_run("El abono podrá actualizarse periódicamente. EL CLIENTE podrá "
                          "darlo de baja en cualquier momento sin penalidad, con un preaviso de 30 (treinta) días.")
        n2r.font.size = Pt(8); n2r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Pie de página con numeración
    foot = doc.sections[0].footer.paragraphs[0]
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("Avanza Digital · avanzadigital.digital · Santa Fe, Argentina · Página ")
    fr.font.size = Pt(7.5); fr.font.color.rgb = RGBColor(0x9A, 0xA3, 0xB2)
    _docx_page_field(foot)

    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()


# ──────────────────────────────────────────────────────────────────────────────
# DISPATCHER + PDF + nombre de archivo
# ──────────────────────────────────────────────────────────────────────────────
def render_contrato_pdf(datos) -> bytes:
    import io
    from xhtml2pdf import pisa

    html_content = render_html(datos)
    buffer = io.BytesIO()
    result = pisa.CreatePDF(io.StringIO(html_content), dest=buffer, encoding="utf-8")
    if result.err:
        raise RuntimeError(f"Error generando PDF con xhtml2pdf: {result.err}")
    return buffer.getvalue()


def render_contrato(datos, formato: str = "pdf") -> bytes:
    """formato: 'pdf' | 'docx'. Devuelve los bytes del contrato."""
    return render_contrato_docx(datos) if str(formato).lower() == "docx" else render_contrato_pdf(datos)


MIME = {"pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}


def nombre_archivo(datos, formato: str = "pdf") -> str:
    base = _strip_accents((datos.cliente_razon_social or "cliente").strip()).lower()
    base = "".join(c if c.isalnum() else "_" for c in base).strip("_") or "cliente"
    ext = "docx" if str(formato).lower() == "docx" else "pdf"
    return f"contrato_avanza_{base}_{datos.plan.replace(' ', '_').lower()}.{ext}"